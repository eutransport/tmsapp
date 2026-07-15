"""File Explorer helpers: tekst-extractie, permissie-checks en validatie.

Alles wat te maken heeft met de nieuwe verkenner-feature (Folder / FileEntry /
FolderPermission) leeft hier, zodat de bestaande signing-services onaangeraakt
blijven.
"""
from __future__ import annotations

import io
import logging
import mimetypes
import os
import re
from typing import Iterable

from django.contrib.auth import get_user_model
from django.core.exceptions import PermissionDenied
from django.db.models import Q

logger = logging.getLogger(__name__)

User = get_user_model()

# --- Bestandsvalidatie ------------------------------------------------------
MAX_UPLOAD_SIZE = 50 * 1024 * 1024  # 50 MB

# Sta gangbare kantoor-/afbeelding-formaten toe. Blokkeer alles wat kan
# worden uitgevoerd (exe, bat, ps1, sh, js, ...).
ALLOWED_EXTENSIONS = {
    # Afbeeldingen
    'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp', 'tiff', 'tif', 'svg', 'heic',
    # PDF
    'pdf',
    # Word / OpenOffice
    'doc', 'docx', 'odt', 'rtf',
    # Excel
    'xls', 'xlsx', 'xlsm', 'ods', 'csv',
    # PowerPoint
    'ppt', 'pptx', 'odp',
    # Platte tekst
    'txt', 'md', 'log', 'json', 'xml',
    # Archieven (worden niet doorzocht, maar wel toegestaan)
    'zip', '7z', 'tar', 'gz',
}

BLOCKED_EXTENSIONS = {
    'exe', 'bat', 'cmd', 'com', 'msi', 'ps1', 'psm1',
    'sh', 'bash', 'zsh', 'php', 'py', 'pyc', 'pl', 'rb',
    'js', 'mjs', 'cjs', 'jsp', 'asp', 'aspx',
    'dll', 'so', 'dylib', 'app', 'apk', 'ipa',
}

_SAFE_NAME_RE = re.compile(r'[^A-Za-z0-9._\-\s()\[\]]+')


def sanitize_filename(name: str) -> str:
    """Verwijder onveilige tekens uit een bestandsnaam en beperk lengte."""
    name = (name or '').strip()
    name = os.path.basename(name)  # verwijder pad-onderdelen
    name = _SAFE_NAME_RE.sub('_', name)
    name = name.strip('. ')
    if not name:
        name = 'bestand'
    return name[:255]


def get_extension(filename: str) -> str:
    if '.' not in filename:
        return ''
    return filename.rsplit('.', 1)[-1].lower()


def validate_upload(uploaded_file) -> tuple[str, str]:
    """Valideer een geüpload bestand. Returnt (safe_name, extension).

    Gooit ``ValueError`` bij een probleem — de view vertaalt dit naar HTTP 400.
    """
    if uploaded_file is None:
        raise ValueError('Geen bestand ontvangen.')
    if uploaded_file.size and uploaded_file.size > MAX_UPLOAD_SIZE:
        raise ValueError(f'Bestand is groter dan {MAX_UPLOAD_SIZE // (1024 * 1024)} MB.')
    safe = sanitize_filename(uploaded_file.name)
    ext = get_extension(safe)
    if ext in BLOCKED_EXTENSIONS:
        raise ValueError('Dit bestandstype is om veiligheidsredenen niet toegestaan.')
    if ext and ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f'Bestandstype ".{ext}" wordt niet ondersteund.')
    return safe, ext


def guess_mime(filename: str) -> str:
    mime, _ = mimetypes.guess_type(filename)
    return mime or 'application/octet-stream'


# --- Permissies -------------------------------------------------------------
def user_is_admin(user) -> bool:
    if not user or not user.is_authenticated:
        return False
    if user.is_superuser:
        return True
    return getattr(user, 'rol', None) == 'admin'


def _folder_permission_qs(user):
    """Alle folder-ids waar deze user expliciet toegang toe heeft."""
    from .models import FolderPermission
    return FolderPermission.objects.filter(user=user).values_list('folder_id', flat=True)


def accessible_folder_ids(user) -> set:
    """Set van folder-ids die deze user mag zien (inclusief afgeleide toegang
    voor sub-mappen van mappen waar iemand expliciet toegang toe heeft).

    Voor admins: geen filtering nodig (roep ``user_is_admin`` apart aan).
    """
    from .models import Folder
    direct_ids = set(_folder_permission_qs(user))
    created_ids = set(Folder.objects.filter(created_by=user).values_list('id', flat=True))
    base = direct_ids | created_ids
    if not base:
        return base
    # Voeg alle nakomelingen toe (BFS). Toegang tot een map = toegang tot
    # onderliggende mappen. Bovenliggende mappen worden NIET automatisch
    # toegankelijk (anders zou je via toegang tot subfolder alsnog een niet-
    # toegankelijke parent kunnen bekijken).
    result = set(base)
    frontier = set(base)
    for _ in range(50):  # veiligheidsgrens
        children = set(
            Folder.objects.filter(parent_id__in=frontier).values_list('id', flat=True)
        )
        new = children - result
        if not new:
            break
        result |= new
        frontier = new
    return result


def user_can_view_folder(user, folder) -> bool:
    if user_is_admin(user):
        return True
    if folder is None:
        return True  # iedereen mag zijn eigen 'root'-zicht hebben
    return folder.id in accessible_folder_ids(user)


def user_can_edit_folder(user, folder) -> bool:
    """Mag deze user in deze map wijzigen (uploaden, submap maken, verwijderen)?"""
    if user_is_admin(user):
        return True
    if folder is None:
        # Root-uploads: alleen admins; anders moet een map gekozen worden.
        return False
    from .models import FolderPermission
    if folder.created_by_id == getattr(user, 'id', None):
        return True
    # Directe toestemming met can_edit=True op deze map of een ancestor.
    ancestors = [folder] + list(folder.get_ancestors())
    ancestor_ids = [a.id for a in ancestors]
    return FolderPermission.objects.filter(
        folder_id__in=ancestor_ids, user=user, can_edit=True
    ).exists()


def filter_folders_for_user(qs, user):
    if user_is_admin(user):
        return qs
    ids = accessible_folder_ids(user)
    return qs.filter(id__in=ids)


def filter_files_for_user(qs, user):
    if user_is_admin(user):
        return qs
    ids = accessible_folder_ids(user)
    # Files in root (folder=NULL) alleen zichtbaar voor uploader.
    return qs.filter(Q(folder_id__in=ids) | Q(folder__isnull=True, uploaded_by=user))


# --- Text extractie ---------------------------------------------------------
MAX_TEXT_CHARS = 500_000  # ~500 KB tekst per bestand in index

def _truncate(text: str) -> str:
    if not text:
        return ''
    if len(text) > MAX_TEXT_CHARS:
        return text[:MAX_TEXT_CHARS]
    return text


def extract_text(uploaded_file, extension: str) -> str:
    """Best-effort tekst-extractie. Faalt stil terug op lege string."""
    extension = (extension or '').lower()
    try:
        # Reset pointer, we hebben mogelijk al de size gelezen.
        if hasattr(uploaded_file, 'seek'):
            try:
                uploaded_file.seek(0)
            except Exception:
                pass
        data = uploaded_file.read()
    except Exception as exc:
        logger.warning('Kon uploaded_file niet lezen voor indexering: %s', exc)
        return ''
    finally:
        try:
            uploaded_file.seek(0)
        except Exception:
            pass

    try:
        if extension == 'pdf':
            return _truncate(_extract_pdf(data))
        if extension == 'docx':
            return _truncate(_extract_docx(data))
        if extension in ('xlsx', 'xlsm'):
            return _truncate(_extract_xlsx(data))
        if extension in ('txt', 'md', 'log', 'csv', 'json', 'xml'):
            try:
                return _truncate(data.decode('utf-8', errors='ignore'))
            except Exception:
                return ''
    except Exception as exc:
        logger.warning('Text-extractie faalde voor .%s: %s', extension, exc)
    return ''


def _extract_pdf(data: bytes) -> str:
    try:
        import fitz  # PyMuPDF
    except Exception:
        return ''
    parts = []
    try:
        with fitz.open(stream=data, filetype='pdf') as doc:
            for page in doc:
                parts.append(page.get_text('text') or '')
    except Exception as exc:
        logger.warning('PDF text-extractie faalde: %s', exc)
    return '\n'.join(parts)


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except Exception:
        return ''
    parts = []
    try:
        doc = Document(io.BytesIO(data))
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
    except Exception as exc:
        logger.warning('DOCX text-extractie faalde: %s', exc)
    return '\n'.join(parts)


def _extract_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except Exception:
        return ''
    parts = []
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if cell is None:
                        continue
                    parts.append(str(cell))
                    if sum(len(p) for p in parts) > MAX_TEXT_CHARS:
                        return ' '.join(parts)
    except Exception as exc:
        logger.warning('XLSX text-extractie faalde: %s', exc)
    return ' '.join(parts)
